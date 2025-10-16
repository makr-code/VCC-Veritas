#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
VERITAS Protected Module
WARNING: This file contains embedded protection keys. 
Modification will be detected and may result in license violations.
"""

# === VERITAS PROTECTION KEYS (DO NOT MODIFY) ===
module_name = "veritas_ui_feedback_system"
module_licenced_organization = "VERITAS_TECH_GMBH"
module_licence_key = "eyJjbGllbnRfaWQi...oNCHqA=="  # Gekuerzt fuer Sicherheit
module_organization_key = "798fc346c8662cb6cb96592a346fd2c4fca287a935ff8c0526f27dcddb62034f"
module_file_key = "899549104f18f5222197cd50ad15dfffaefc8975376f9e1597972b0d48bf044f"
module_version = "1.0"
module_protection_level = 1
# === END PROTECTION KEYS ===
"""
Fe        # Wiederholen-Button
        self._create_action_button(button_frame, "🔄", "Erneut fragen", self._repeat_question)
        
        self._add_spacer(button_frame, 15)
        
        # Feedback-Bereich
        feedback_frame = tk.Frame(button_frame)
        feedback_frame.pack(side="left")
        
        # Feedback-Buttons (vergrößert)
        self.btn_like = self._create_feedback_button(feedback_frame, "👍", True, "Antwort positiv bewerten", font_size=14)
        self.btn_dislike = self._create_feedback_button(feedback_frame, "👎", False, "Antwort negativ bewerten", font_size=14)m für Veritas Chat
Basiert auf dem ursprünglichen Design mit modernen Verbesserungen
"""

import tkinter as tk
from tkinter import ttk, messagebox
import datetime
import logging
import json
from typing import Dict, List, Optional, Callable
from frontend.ui.veritas_ui_components import Tooltip

logger = logging.getLogger(__name__)


class MessageFeedbackWidget:
    def export_feedback_to_csv_dialog(self):
        """Bietet einen 'Speichern unter...' Dialog für den CSV-Export."""
        from tkinter import filedialog
        file_path = filedialog.asksaveasfilename(
            title="Feedback als CSV speichern",
            defaultextension=".csv",
            filetypes=[("CSV-Datei", "*.csv"), ("Alle Dateien", "*.*")]
        )
        if file_path:
            self.export_feedback_to_csv(file_path)

    def export_feedback_to_pdf_dialog(self):
        """Bietet einen 'Speichern unter...' Dialog für den PDF-Export."""
        from tkinter import filedialog
        file_path = filedialog.asksaveasfilename(
            title="Feedback als PDF speichern",
            defaultextension=".pdf",
            filetypes=[("PDF-Datei", "*.pdf"), ("Alle Dateien", "*.*")]
        )
        if file_path:
            self.export_feedback_to_pdf(file_path)

    def export_feedback_to_json_dialog(self):
        """Bietet einen 'Speichern unter...' Dialog für den JSON-Export."""
        from tkinter import filedialog
        file_path = filedialog.asksaveasfilename(
            title="Feedback als JSON speichern",
            defaultextension=".json",
            filetypes=[("JSON-Datei", "*.json"), ("Alle Dateien", "*.*")]
        )
        if file_path:
            self.export_feedback_to_json(file_path)

    def export_feedback_to_word_dialog(self):
        """Bietet einen 'Speichern unter...' Dialog für den Word-Export."""
        from tkinter import filedialog
        file_path = filedialog.asksaveasfilename(
            title="Feedback als Word speichern",
            defaultextension=".docx",
            filetypes=[("Word-Datei", "*.docx"), ("Alle Dateien", "*.*")]
        )
        if file_path:
            self.export_feedback_to_word(file_path)
    def export_feedback_to_csv(self, filename="feedback_export.csv"):
        """Exportiert das Feedback als CSV-Datei."""
        import csv
        try:
            with open(filename, 'w', newline='', encoding='utf-8') as csvfile:
                writer = csv.writer(csvfile)
                writer.writerow(["Bewertung", "Kommentar", "Zeit"])
                for fb in self.feedback_list:
                    writer.writerow([
                        fb.get('rating', ''),
                        fb.get('detailed_feedback', ''),
                        fb.get('timestamp', '')
                    ])
            logger.info(f"Feedback als CSV gespeichert: {filename}")
        except Exception as e:
            logger.error(f"Fehler beim CSV-Export: {e}")

    def export_feedback_to_pdf(self, filename="feedback_export.pdf"):
        """Exportiert das Feedback als PDF-Datei."""
        try:
            from fpdf import FPDF
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            pdf.cell(200, 10, txt="Feedback Export", ln=True, align='C')
            for fb in self.feedback_list:
                pdf.cell(200, 10, txt=f"Bewertung: {fb.get('rating', '')}", ln=True)
                pdf.cell(200, 10, txt=f"Kommentar: {fb.get('detailed_feedback', '')}", ln=True)
                pdf.cell(200, 10, txt=f"Zeit: {fb.get('timestamp', '')}", ln=True)
                pdf.cell(200, 10, txt="---", ln=True)
            pdf.output(filename)
            logger.info(f"Feedback als PDF gespeichert: {filename}")
        except Exception as e:
            logger.error(f"Fehler beim PDF-Export: {e}")

    def export_feedback_to_json(self, filename="feedback_export.json"):
        """Exportiert das Feedback als JSON-Datei."""
        import json
        try:
            with open(filename, 'w', encoding='utf-8') as f:
                json.dump(self.feedback_list, f, ensure_ascii=False, indent=2)
            logger.info(f"Feedback als JSON gespeichert: {filename}")
        except Exception as e:
            logger.error(f"Fehler beim JSON-Export: {e}")
    
    def copy_feedback_to_clipboard(self):
        """Kopiert das letzte Feedback als Text in die Zwischenablage."""
        if self.feedback_list:
            last_feedback = self.feedback_list[-1]
            text = f"Bewertung: {last_feedback.get('rating', '')}\nKommentar: {last_feedback.get('detailed_feedback', '')}"
            try:
                import pyperclip
                pyperclip.copy(text)
                logger.info("Feedback in Zwischenablage kopiert.")
            except Exception as e:
                logger.error(f"Fehler beim Kopieren des Feedbacks: {e}")

    def export_feedback_to_word(self, filename="feedback_export.docx"):
        """Exportiert das Feedback als Word-Datei."""
        try:
            from docx import Document
            doc = Document()
            doc.add_heading("Feedback Export", 0)
            for fb in self.feedback_list:
                doc.add_paragraph(f"Bewertung: {fb.get('rating', '')}")
                doc.add_paragraph(f"Kommentar: {fb.get('detailed_feedback', '')}")
                doc.add_paragraph(f"Zeit: {fb.get('timestamp', '')}")
                doc.add_paragraph("---")
            doc.save(filename)
            logger.info(f"Feedback als Word-Datei gespeichert: {filename}")
        except Exception as e:
            logger.error(f"Fehler beim Word-Export: {e}")
    """
    Widget für Feedback-Buttons und -Funktionen einer einzelnen Nachricht.
    """
    
    def __init__(self, parent_display, message_data: Dict, save_callback: Optional[Callable] = None):
        """
        Initialisiert das Feedback-Widget.
        
        Args:
            parent_display: Das Text-Widget wo die Nachricht angezeigt wird
            message_data: Die Nachrichtendaten mit Metadaten
            save_callback: Callback zum Speichern von Änderungen
        """
        self.parent_display = parent_display
        self.message_data = message_data
        self.save_callback = save_callback
        
        self.current_feedback_rating = None
        self.feedback_buttons = {}
        self.feedback_input_frame = None
        self.feedback_entry = None
        
        # Feedback aus Metadaten laden
        self.feedback_list = message_data.get("metadata", {}).get("user_feedback", [])
        
    def create_feedback_buttons(self):
        """Erstellt die Feedback-Button-Leiste."""
        self.parent_display.config(state='normal')
        self.parent_display.insert(tk.END, "\n")

        # Container für alle Buttons
        button_container = tk.Frame(self.parent_display, relief="flat", bd=0)
        
        # Button-Frame
        button_frame = tk.Frame(button_container)
        button_frame.pack(padx=8, pady=6)

        # Wiederholen-Button
        self._create_action_button(button_frame, "🔄", "Erneut fragen", self._repeat_question)
        
        self._add_spacer(button_frame, 15)
        
        # Feedback-Bereich
        feedback_frame = tk.Frame(button_frame)
        feedback_frame.pack(side="left")
        
        # Feedback-Buttons (vergrößert)
        self.btn_like = self._create_feedback_button(feedback_frame, "👍", True, "Antwort positiv bewerten")
        self.btn_dislike = self._create_feedback_button(feedback_frame, "👎", False, "Antwort negativ bewerten")      
        # Feedback-Eingabefeld (versteckt)
        self.feedback_input_frame = tk.Frame(feedback_frame)
        
        self.feedback_entry = tk.Entry(
            self.feedback_input_frame,
            bd=1, relief="solid", width=25,
            highlightthickness=0
        )
        self.feedback_entry.pack(side="left", padx=(5, 2))
        
        self.feedback_send_btn = tk.Button(
            self.feedback_input_frame, text="⏎",
            relief="flat", bd=0, padx=4, pady=1,
            command=self._send_detailed_feedback,
            cursor="hand2"
        )
        self.feedback_send_btn.pack(side="left")
        
        self._add_spacer(button_frame, 12)
        
        # Aktions-Buttons
        actions_frame = tk.Frame(button_frame)
        actions_frame.pack(side="left")
        
        self._create_action_button(actions_frame, "🔊", "Antwort vorlesen", self._speak_answer, font_size=14)
        self._create_action_button(actions_frame, "🔖", "Antwort merken (Bookmark)", self._bookmark_answer, font_size=14)  
        self._create_action_button(actions_frame, "📋", "In Zwischenablage", self._copy_answer, font_size=14)
        self._create_action_button(actions_frame, "💾", "Antwort speichern unter...", self._show_save_dialog, font_size=14)
        self._create_action_button(actions_frame, "🔗", "Antwort teilen", self._share_answer, font_size=14)

    def _bookmark_answer(self):
        """Merkt sich die Antwort als Bookmark (z.B. speichert sie in einer Liste oder Datei)."""
        try:
            content = self.message_data.get("content", "")
            if content:
                # Beispiel: Bookmark lokal speichern (Datei oder Attribut)
                bookmark = {
                    "content": content,
                    "timestamp": datetime.datetime.now().isoformat()
                }
                # Hier als Datei speichern, alternativ in einer globalen/bookmark-Liste
                with open("bookmarks.txt", "a", encoding="utf-8") as f:
                    f.write(json.dumps(bookmark, ensure_ascii=False) + "\n")
                messagebox.showinfo("Bookmark", "Antwort wurde als Bookmark gespeichert.")
            else:
                messagebox.showwarning("Bookmark", "Keine Antwort zum Merken vorhanden.")
        except Exception as e:
            logger.error(f"Fehler beim Bookmark: {e}")
            messagebox.showerror("Fehler", f"Fehler beim Bookmark:\n{e}")

    def _share_answer(self):
        """Teilt die Antwort (z.B. kopiert einen Link oder öffnet einen Dialog)."""
        try:
            # Beispiel: Kopiert die Antwort als Text in die Zwischenablage und zeigt eine Info
            import pyperclip
            content = self.message_data.get("content", "")
            if content:
                pyperclip.copy(content)
                messagebox.showinfo("Teilen", "Antwort wurde in die Zwischenablage kopiert und kann geteilt werden.")
            else:
                messagebox.showwarning("Teilen", "Keine Antwort zum Teilen vorhanden.")
        except Exception as e:
            logger.error(f"Fehler beim Teilen: {e}")
            messagebox.showerror("Fehler", f"Fehler beim Teilen:\n{e}")

    def _show_save_dialog(self):
        """Zeigt einen Dialog zur Auswahl des Exportformats und ruft den passenden 'Speichern unter...' Dialog auf."""
        from tkinter import simpledialog
        formats = [
            ("CSV", self.export_feedback_to_csv_dialog),
            ("PDF", self.export_feedback_to_pdf_dialog),
            ("JSON", self.export_feedback_to_json_dialog),
            ("Word", self.export_feedback_to_word_dialog)
        ]
        format_names = [f[0] for f in formats]
        # Einfacher Auswahl-Dialog
        selected = simpledialog.askstring(
            "Exportformat wählen",
            "Bitte Exportformat eingeben (CSV, PDF, JSON, Word):",
            initialvalue="CSV"
        )
        if selected:
            selected = selected.strip().upper()
            for name, func in formats:
                if name.upper() == selected:
                    func()
                    break
        
        # Feedback-Status anzeigen
        self._add_feedback_status_display(button_container)
        self._update_feedback_button_states()

        # Widget in Display einbetten
        self.parent_display.window_create(tk.END, window=button_container)
        self.parent_display.insert(tk.END, "\n\n")
        self.parent_display.config(state='disabled')
        
        return button_container
    
    def _create_action_button(self, parent, text: str, tooltip: str, command: Callable, font_size: int = 12):
        """Erstellt einen Aktions-Button mit einheitlichem Styling."""
        btn = tk.Button(
            parent, text=text,
            relief="flat", bd=0,
            padx=6, pady=4,
            command=command,
            cursor="hand2",
            font=("Segoe UI Emoji", font_size)
        )
        
        # Hover-Effekt
        def on_enter(e):
            btn.config(bg="#e3f2fd")
        def on_leave(e):
            btn.config(bg="#f5f5f5")
            
        btn.bind("<Enter>", on_enter)
        btn.bind("<Leave>", on_leave)
        btn.pack(side="left", padx=2)
        
        # Tooltip mit der richtigen Tooltip-Klasse
        Tooltip(btn, tooltip)
        
        return btn
    
    def _create_feedback_button(self, parent, text: str, is_positive: bool, tooltip: str):
        """Erstellt einen Feedback-Button mit vergrößerter Schrift."""
        btn = tk.Button(
            parent, text=text,
            relief="flat", bd=0,
            padx=8, pady=6,
            font=("Segoe UI Emoji", 14),
            command=lambda: self._toggle_feedback(is_positive),
            cursor="hand2"
        )
        
        # Hover-Effekt nur wenn nicht aktiv
        def on_enter(e):
            if btn.cget('bg') == "":
                btn.config()
        def on_leave(e):
            if btn.cget('bg') == "":
                btn.config()
                
        btn.bind("<Enter>", on_enter)
        btn.bind("<Leave>", on_leave)
        btn.pack(side="left", padx=1)
        
        # Tooltip mit der richtigen Tooltip-Klasse
        Tooltip(btn, tooltip)
        
        self.feedback_buttons[is_positive] = btn
        return btn
    
    def _add_spacer(self, parent, width: int):
        """Fügt einen Abstandshalter hinzu."""
        spacer = tk.Frame(parent, width=width)
        spacer.pack(side="left")
    
    def _toggle_feedback(self, is_positive: bool):
        """Aktiviert/Deaktiviert Feedback und zeigt Eingabefeld."""
        # Farben für Feedback-Status
        positive_color = "#90EE90"  # Helles Grün
        negative_color = "#FFB6C1"  # Helles Rot
        neutral_color = "#F0F0F0"   # Standard
        
        # Reset beide Buttons
        self.feedback_buttons[True].config(bg=neutral_color)
        self.feedback_buttons[False].config(bg=neutral_color)
        
        # Aktueller Button aktivieren
        self.feedback_buttons[is_positive].config(
            bg=positive_color if is_positive else negative_color
        )
        
        self.current_feedback_rating = "positive" if is_positive else "negative"
        
        # Eingabefeld anzeigen
        self.feedback_input_frame.pack(side="left", padx=(8, 0))
        self.feedback_entry.focus_set()
        
        # Placeholder-Text
        placeholder = "Kommentar (optional)..."
        self.feedback_entry.delete(0, tk.END)
        self.feedback_entry.insert(0, placeholder)
        self.feedback_entry.config()
        
        # Placeholder-Verhalten
        def on_focus_in(event):
            if self.feedback_entry.get() == placeholder:
                self.feedback_entry.delete(0, tk.END)
                self.feedback_entry.config()
        
        def on_focus_out(event):
            if not self.feedback_entry.get().strip():
                self.feedback_entry.insert(0, placeholder)
                self.feedback_entry.config()
        
        # Event-Bindings
        self.feedback_entry.bind("<FocusIn>", on_focus_in)
        self.feedback_entry.bind("<FocusOut>", on_focus_out)
        self.feedback_entry.bind("<Return>", lambda e: self._send_detailed_feedback())
    
    def _send_detailed_feedback(self):
        """Sendet das Feedback mit Detail-Text."""
        if not self.current_feedback_rating:
            return
            
        # Feedback-Text extrahieren
        feedback_text = self.feedback_entry.get().strip()
        if feedback_text == "Kommentar (optional)...":
            feedback_text = ""
        
        # Feedback-Daten erstellen
        feedback_data = {
            "rating": self.current_feedback_rating,
            "detailed_feedback": feedback_text,
            "timestamp": datetime.datetime.now().isoformat(),
            "user_id": "current_user",
            "feedback_version": "2.0"
        }
        
        # Zu Metadaten hinzufügen
        if "metadata" not in self.message_data:
            self.message_data["metadata"] = {}
            
        if "user_feedback" not in self.message_data["metadata"]:
            self.message_data["metadata"]["user_feedback"] = []
            
        self.message_data["metadata"]["user_feedback"].append(feedback_data)
        self.feedback_list = self.message_data["metadata"]["user_feedback"]
        
        # Speichern falls Callback vorhanden
        if self.save_callback:
            try:
                self.save_callback()
            except Exception as e:
                logger.error(f"Fehler beim Speichern des Feedbacks: {e}")
        
        # UI-Updates
        self.feedback_input_frame.pack_forget()
        self._update_feedback_status_display()
        
        # Bestätigung
        logger.info(f"Feedback gespeichert: {self.current_feedback_rating} - {feedback_text}")
    
    def _add_feedback_status_display(self, parent_container):
        """Zeigt vorhandenes Feedback an."""
        if not self.feedback_list:
            return
            
        # Status-Frame
        status_frame = tk.Frame(parent_container)
        status_frame.pack(fill="x", padx=8, pady=(3, 0))
        
        # Letztes Feedback
        last_feedback = self.feedback_list[-1]
        rating = last_feedback.get("rating", "unknown")
        icon = "👍" if rating == "positive" else "👎" if rating == "negative" else "❓"
        
        feedback_text = last_feedback.get("detailed_feedback", "").strip()
        if feedback_text:
            display_text = feedback_text[:30] + "..." if len(feedback_text) > 30 else feedback_text
        else:
            display_text = "Bewertet"
        
        # Status-Label
        status_label = tk.Label(
            status_frame,
            text=f"{icon} {display_text}"
        )
        status_label.pack(anchor="w", padx=2, pady=1)
        status_label.config(cursor="hand2")  # Cursor separat setzen
        
        # Click-Handler für Details
        status_label.bind("<Button-1>", lambda e: self._show_feedback_details())
        
        self.feedback_status_label = status_label
    
    def _update_feedback_status_display(self):
        """Aktualisiert die Feedback-Status-Anzeige."""
        if hasattr(self, 'feedback_status_label') and self.feedback_list:
            last_feedback = self.feedback_list[-1]
            rating = last_feedback.get("rating", "unknown")
            icon = "👍" if rating == "positive" else "👎" if rating == "negative" else "❓"
            
            feedback_text = last_feedback.get("detailed_feedback", "").strip()
            if feedback_text:
                display_text = feedback_text[:30] + "..." if len(feedback_text) > 30 else feedback_text
            else:
                display_text = "Bewertet"
                
            self.feedback_status_label.config(text=f"{icon} {display_text}")
    
    def _update_feedback_button_states(self):
        """Aktualisiert die Button-Zustände basierend auf vorhandenem Feedback."""
        if not self.feedback_list:
            return
            
        # Letztes Feedback anzeigen
        last_feedback = self.feedback_list[-1]
        rating = last_feedback.get("rating")
        
        if rating == "positive":
            self.feedback_buttons[True].config()
        elif rating == "negative":
            self.feedback_buttons[False].config()
    
    def _show_feedback_details(self):
        """Zeigt alle Feedbacks in einem Dialog."""
        if not self.feedback_list:
            return
            
        # Dialog erstellen
        dialog = tk.Toplevel()
        dialog.title("Feedback-Verlauf")
        dialog.geometry("500x400")
        dialog.transient(self.parent_display.master)
        dialog.grab_set()
        
        # Header
        header_frame = tk.Frame(dialog, pady=10)
        header_frame.pack(fill="x")
        
        title_label = tk.Label(
            header_frame,
            text=f"📝 Feedback-Verlauf ({len(self.feedback_list)} Einträge)")
        title_label.pack()
        
        # Scrollbare Liste
        main_frame = tk.Frame(dialog)
        main_frame.pack(fill="both", expand=True, padx=10, pady=10)
        
        canvas = tk.Canvas(main_frame, highlightthickness=0)
        scrollbar = ttk.Scrollbar(main_frame, orient="vertical", command=canvas.yview)
        scrollable_frame = tk.Frame(canvas)
        
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        # Feedback-Einträge
        for i, feedback in enumerate(self.feedback_list, 1):
            self._create_feedback_item(scrollable_frame, feedback, i)
        
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        # Schließen-Button
        close_btn = tk.Button(
            dialog, text="Schließen",
            relief="flat", bd=0, padx=20, pady=8,
            command=dialog.destroy
        )
        close_btn.pack(pady=10)
    
    def _create_feedback_item(self, parent, feedback: Dict, index: int):
        """Erstellt ein Feedback-Item im Dialog."""
        item_frame = tk.Frame(parent, relief="solid", bd=1)
        item_frame.pack(fill="x", pady=3, padx=5)
        
        # Header
        header_frame = tk.Frame(item_frame)
        header_frame.pack(fill="x", padx=10, pady=(8, 5))
        
        rating = feedback.get("rating", "unknown")
        icon = "👍" if rating == "positive" else "👎" if rating == "negative" else "❓"
        rating_text = "Positiv" if rating == "positive" else "Negativ" if rating == "negative" else "Unbekannt"
        
        header_label = tk.Label(
            header_frame,
            text=f"#{index} {icon} {rating_text}")
        header_label.pack(side="left")
        
        # Zeitstempel
        timestamp = feedback.get("timestamp", "")
        if timestamp:
            try:
                dt = datetime.datetime.fromisoformat(timestamp.replace('Z', '+00:00'))
                time_str = dt.strftime("%d.%m.%Y %H:%M")
            except:
                time_str = timestamp[:16]
                
            time_label = tk.Label(
                header_frame,
                text=time_str)
            time_label.pack(side="right")
        
        # Detail-Text
        detail_text = feedback.get("detailed_feedback", "").strip()
        if detail_text:
            detail_label = tk.Label(
                item_frame,
                text=detail_text,
                anchor="w", justify="left",
                wraplength=450
            )
            detail_label.pack(fill="x", padx=10, pady=(0, 8))
    
    # Aktions-Methoden (Stubs - können von Chat-Fenster implementiert werden)
    def _repeat_question(self):
        """Wiederholt die ursprüngliche Frage."""
        logger.info("Wiederholen-Funktion aufgerufen")
    
    def _speak_answer(self):
        """Liest die Antwort vor."""
        logger.info("Vorlese-Funktion aufgerufen")
    
    def _copy_answer(self):
        """Kopiert die Antwort in die Zwischenablage."""
        logger.info("Kopieren-Funktion aufgerufen")
    
    def _save_answer(self):
        """Speichert die Antwort in eine Datei."""
        logger.info("Speichern-Funktion aufgerufen")


class FeedbackManager:
    def export_all_feedback_to_word(self, filename="all_feedback_export.docx"):
        """Exportiert das gesamte Feedback aller Nachrichten als Word-Datei."""
        try:
            from docx import Document
            doc = Document()
            doc.add_heading("Gesamtes Feedback Export", 0)
            for msg_id, widget in self.feedback_widgets.items():
                doc.add_heading(f"Nachricht: {msg_id}", level=1)
                for fb in widget.feedback_list:
                    doc.add_paragraph(f"Bewertung: {fb.get('rating', '')}")
                    doc.add_paragraph(f"Kommentar: {fb.get('detailed_feedback', '')}")
                    doc.add_paragraph(f"Zeit: {fb.get('timestamp', '')}")
                    doc.add_paragraph("---")
            doc.save(filename)
            logger.info(f"Gesamtes Feedback als Word-Datei gespeichert: {filename}")
        except Exception as e:
            logger.error(f"Fehler beim Word-Export: {e}")
    """
    Manager für das gesamte Feedback-System eines Chat-Fensters.
    """
    
    def __init__(self, chat_window):
        """
        Initialisiert den Feedback-Manager.
        
        Args:
            chat_window: Das Chat-Fenster
        """
        self.chat_window = chat_window
        self.feedback_widgets = {}  # message_id -> MessageFeedbackWidget
    
    def add_message_feedback(self, message_id: str, message_data: Dict):
        """
        Fügt Feedback-Funktionen zu einer Nachricht hinzu.
        
        Args:
            message_id: Eindeutige ID der Nachricht
            message_data: Nachrichtendaten
        """
        if message_data.get("role") != "assistant":
            return  # Nur für Assistant-Nachrichten
            
        # Feedback-Widget erstellen
        feedback_widget = MessageFeedbackWidget(
            parent_display=self.chat_window.chat_display,
            message_data=message_data,
            save_callback=self.chat_window.save_chat
        )
        
        # Aktions-Callbacks setzen
        feedback_widget._repeat_question = self._repeat_last_question
        feedback_widget._speak_answer = lambda: self._speak_answer(message_data)
        feedback_widget._copy_answer = lambda: self._copy_answer(message_data) 
        feedback_widget._save_answer = lambda: self._save_answer(message_data)
        
        # Widget speichern und Button-Leiste erstellen
        self.feedback_widgets[message_id] = feedback_widget
        feedback_widget.create_feedback_buttons()
    
    def _repeat_last_question(self):
        """Wiederholt die letzte Benutzerfrage."""
        try:
            if hasattr(self.chat_window, 'messages') and self.chat_window.messages:
                # Finde letzte User-Nachricht
                for message in reversed(self.chat_window.messages):
                    if message.get("role") == "user":
                        # Setze in Eingabefeld und sende
                        content = message.get("content", "")
                        self.chat_window.entry_field.delete(0, tk.END)
                        self.chat_window.entry_field.insert(0, content)
                        self.chat_window._send_message()
                        break
        except Exception as e:
            logger.error(f"Fehler beim Wiederholen der Frage: {e}")
    
    def _speak_answer(self, message_data: Dict):
        """Liest eine Antwort vor."""
        try:
            content = message_data.get("content", "")
            if content:
                # Hier könnte TTS-Integration erfolgen
                messagebox.showinfo("TTS", f"Vorlesen: {content[:50]}...")
        except Exception as e:
            logger.error(f"Fehler beim Vorlesen: {e}")
    
    def _copy_answer(self, message_data: Dict):
        """Kopiert eine Antwort in die Zwischenablage."""
        try:
            content = message_data.get("content", "")
            if content:
                self.chat_window.clipboard_clear()
                self.chat_window.clipboard_append(content)
                self.chat_window.update_status("Antwort in Zwischenablage kopiert", "success")
        except Exception as e:
            logger.error(f"Fehler beim Kopieren: {e}")
    
    def _save_answer(self, message_data: Dict):
        """Speichert eine Antwort in eine Datei."""
        try:
            from tkinter import filedialog
            content = message_data.get("content", "")
            if content:
                file_path = filedialog.asksaveasfilename(
                    title="Antwort speichern",
                    defaultextension=".txt",
                    filetypes=[("Textdatei", "*.txt"), ("Alle Dateien", "*.*")]
                )
                if file_path:
                    with open(file_path, 'w', encoding='utf-8') as f:
                        f.write(content)
                    self.chat_window.update_status(f"Antwort gespeichert: {file_path}", "success")
        except Exception as e:
            logger.error(f"Fehler beim Speichern: {e}")
            messagebox.showerror("Fehler", f"Fehler beim Speichern:\n{e}")
