#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
VERITAS Office Export Service
Export Chat-Konversationen als Word/Excel-Dokumente
"""

import logging
from datetime import datetime, timedelta
from pathlib import Path
from typing import List, Dict, Optional, Any
import re

# Word Export
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("⚠️ python-docx nicht verfügbar")

# Excel Export
try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    from openpyxl.utils import get_column_letter
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False
    logging.warning("⚠️ openpyxl nicht verfügbar")

logger = logging.getLogger(__name__)


class OfficeExportService:
    """
    Service für Office-Dokument-Export
    
    Features:
    - Word-Export (.docx) mit Formatierung
    - Excel-Export (.xlsx) mit Statistiken
    - Zeitraum-Filter
    - Markdown → Word Konvertierung
    """
    
    def __init__(self, output_dir: Optional[Path] = None):
        """
        Initialisiert Export-Service
        
        Args:
            output_dir: Ausgabe-Verzeichnis (default: ./exports/)
        """
        self.output_dir = output_dir or Path("./exports")
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        logger.info(f"✅ OfficeExportService initialisiert: {self.output_dir}")
    
    # ===== WORD EXPORT =====
    
    def export_to_word(
        self,
        chat_messages: List[Dict],
        filename: Optional[str] = None,
        title: str = "VERITAS Chat-Protokoll",
        include_metadata: bool = True,
        include_sources: bool = True
    ) -> Path:
        """
        Exportiert Chat als Word-Dokument
        
        Args:
            chat_messages: Liste von Chat-Messages
            filename: Dateiname (default: auto-generiert)
            title: Dokument-Titel
            include_metadata: Metriken einschließen
            include_sources: Quellen einschließen
            
        Returns:
            Path zum exportierten Dokument
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx nicht installiert. Install: pip install python-docx")
        
        # Create Document
        doc = Document()
        
        # Setup Styles
        self._setup_word_styles(doc)
        
        # Title
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata Header
        doc.add_paragraph(
            f"Exportiert am: {datetime.now().strftime('%d.%m.%Y %H:%M')}",
            style='Subtitle'
        )
        doc.add_paragraph(
            f"Anzahl Messages: {len(chat_messages)}",
            style='Subtitle'
        )
        doc.add_paragraph("")  # Spacing
        
        # Messages
        for i, msg in enumerate(chat_messages, 1):
            self._add_message_to_word(
                doc, msg, i,
                include_metadata=include_metadata,
                include_sources=include_sources
            )
        
        # Footer
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"Generiert von VERITAS v3.17.0 | {datetime.now().strftime('%d.%m.%Y')}"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Save
        if not filename:
            filename = f"veritas_chat_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        output_path = self.output_dir / filename
        doc.save(str(output_path))
        
        logger.info(f"✅ Word-Export erfolgreich: {output_path}")
        return output_path
    
    def _setup_word_styles(self, doc: Document):
        """Konfiguriert Word-Styles"""
        styles = doc.styles
        
        # User Message Style
        try:
            user_style = styles.add_style('UserMessage', WD_STYLE_TYPE.PARAGRAPH)
            user_style.font.name = 'Segoe UI'
            user_style.font.size = Pt(11)
            user_style.font.color.rgb = RGBColor(0, 102, 204)  # Blue
            user_style.paragraph_format.space_before = Pt(6)
            user_style.paragraph_format.space_after = Pt(3)
        except ValueError:
            pass  # Style already exists
        
        # Assistant Message Style
        try:
            assistant_style = styles.add_style('AssistantMessage', WD_STYLE_TYPE.PARAGRAPH)
            assistant_style.font.name = 'Segoe UI'
            assistant_style.font.size = Pt(11)
            assistant_style.paragraph_format.space_before = Pt(3)
            assistant_style.paragraph_format.space_after = Pt(6)
        except ValueError:
            pass
    
    def _add_message_to_word(
        self,
        doc: Document,
        message: Dict,
        index: int,
        include_metadata: bool,
        include_sources: bool
    ):
        """Fügt Message zu Word-Dokument hinzu"""
        role = message.get('role', 'user')
        content = message.get('content', '')
        timestamp = message.get('timestamp', '')
        
        # Header
        if role == 'user':
            header = doc.add_paragraph(f"🙋 User ({index})", style='UserMessage')
            header.bold = True
        else:
            header = doc.add_paragraph(f"🤖 Assistant ({index})", style='AssistantMessage')
            header.bold = True
        
        # Timestamp
        if timestamp:
            ts_para = doc.add_paragraph(f"⏰ {timestamp}", style='Body Text')
            ts_para.runs[0].font.size = Pt(9)
            ts_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)
        
        # Content (Markdown → Plain Text mit Formatierung)
        content_paras = self._markdown_to_word(doc, content)
        
        # Metadata
        if include_metadata and role == 'assistant':
            metadata = message.get('metadata', {})
            if metadata:
                meta_para = doc.add_paragraph(style='Body Text')
                meta_para.add_run("📊 Metriken: ").bold = True
                
                metrics = []
                if metadata.get('confidence'):
                    metrics.append(f"Confidence: {metadata['confidence']}%")
                if metadata.get('duration'):
                    metrics.append(f"Dauer: {metadata['duration']:.1f}s")
                if metadata.get('sources_count'):
                    metrics.append(f"Quellen: {metadata['sources_count']}")
                
                meta_para.add_run(" | ".join(metrics))
        
        # Sources
        if include_sources and role == 'assistant':
            sources = message.get('sources', [])
            if sources:
                doc.add_paragraph("📚 Quellen:", style='Heading 3')
                for source in sources:
                    source_para = doc.add_paragraph(style='List Bullet')
                    source_para.add_run(source.get('title', 'Unbekannt')).bold = True
                    if source.get('page'):
                        source_para.add_run(f" (Seite {source['page']})")
        
        # Separator
        doc.add_paragraph("─" * 80, style='Body Text')
    
    def _markdown_to_word(self, doc: Document, markdown_text: str) -> List:
        """
        Konvertiert Markdown zu Word-Paragraphen (basic)
        
        Args:
            doc: Document
            markdown_text: Markdown-Text
            
        Returns:
            Liste von erstellten Paragraphen
        """
        paragraphs = []
        
        # Split by lines
        lines = markdown_text.split('\n')
        
        for line in lines:
            line = line.strip()
            
            if not line:
                continue
            
            # Heading 1
            if line.startswith('# '):
                p = doc.add_paragraph(line[2:], style='Heading 1')
                paragraphs.append(p)
            
            # Heading 2
            elif line.startswith('## '):
                p = doc.add_paragraph(line[3:], style='Heading 2')
                paragraphs.append(p)
            
            # Heading 3
            elif line.startswith('### '):
                p = doc.add_paragraph(line[4:], style='Heading 3')
                paragraphs.append(p)
            
            # List
            elif line.startswith('- ') or line.startswith('* '):
                p = doc.add_paragraph(line[2:], style='List Bullet')
                paragraphs.append(p)
            
            # Numbered List
            elif re.match(r'^\d+\.\s', line):
                content = re.sub(r'^\d+\.\s', '', line)
                p = doc.add_paragraph(content, style='List Number')
                paragraphs.append(p)
            
            # Code Block
            elif line.startswith('```'):
                continue  # Skip code fence
            
            # Bold/Italic (simple)
            else:
                p = doc.add_paragraph(style='Body Text')
                
                # Replace **bold**
                parts = re.split(r'\*\*(.*?)\*\*', line)
                for i, part in enumerate(parts):
                    if i % 2 == 0:
                        p.add_run(part)
                    else:
                        p.add_run(part).bold = True
                
                paragraphs.append(p)
        
        return paragraphs
    
    # ===== EXCEL EXPORT =====
    
    def export_to_excel(
        self,
        chat_messages: List[Dict],
        feedback_stats: Optional[Dict] = None,
        filename: Optional[str] = None
    ) -> Path:
        """
        Exportiert Chat als Excel-Workbook
        
        Args:
            chat_messages: Liste von Chat-Messages
            feedback_stats: Feedback-Statistiken (optional)
            filename: Dateiname (default: auto-generiert)
            
        Returns:
            Path zum exportierten Dokument
        """
        if not OPENPYXL_AVAILABLE:
            raise ImportError("openpyxl nicht installiert. Install: pip install openpyxl")
        
        # Create Workbook
        wb = Workbook()
        
        # Sheet 1: Messages
        ws_messages = wb.active
        ws_messages.title = "Chat Messages"
        self._create_messages_sheet(ws_messages, chat_messages)
        
        # Sheet 2: Statistics
        ws_stats = wb.create_sheet("Statistiken")
        self._create_statistics_sheet(ws_stats, chat_messages, feedback_stats)
        
        # Sheet 3: Sources
        ws_sources = wb.create_sheet("Quellen")
        self._create_sources_sheet(ws_sources, chat_messages)
        
        # Save
        if not filename:
            filename = f"veritas_chat_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        
        output_path = self.output_dir / filename
        wb.save(str(output_path))
        
        logger.info(f"✅ Excel-Export erfolgreich: {output_path}")
        return output_path
    
    def _create_messages_sheet(self, ws, messages: List[Dict]):
        """Erstellt Messages-Sheet"""
        # Header
        headers = ['Nr.', 'Role', 'Timestamp', 'Content', 'Confidence', 'Duration', 'Sources']
        ws.append(headers)
        
        # Style Header
        for cell in ws[1]:
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
            cell.alignment = Alignment(horizontal="center", vertical="center")
        
        # Data
        for i, msg in enumerate(messages, 1):
            role = msg.get('role', 'user')
            timestamp = msg.get('timestamp', '')
            content = msg.get('content', '')[:200]  # Max 200 chars
            metadata = msg.get('metadata', {})
            
            row = [
                i,
                role.upper(),
                timestamp,
                content,
                metadata.get('confidence', ''),
                metadata.get('duration', ''),
                len(msg.get('sources', []))
            ]
            ws.append(row)
        
        # Column widths
        ws.column_dimensions['A'].width = 8
        ws.column_dimensions['B'].width = 12
        ws.column_dimensions['C'].width = 20
        ws.column_dimensions['D'].width = 60
        ws.column_dimensions['E'].width = 12
        ws.column_dimensions['F'].width = 12
        ws.column_dimensions['G'].width = 10
        
        # Freeze header
        ws.freeze_panes = 'A2'
    
    def _create_statistics_sheet(self, ws, messages: List[Dict], feedback_stats: Optional[Dict]):
        """Erstellt Statistiken-Sheet"""
        ws.append(['VERITAS Chat-Statistiken'])
        ws['A1'].font = Font(size=16, bold=True)
        ws.append([])
        
        # Message Statistics
        ws.append(['Nachrichten-Statistik'])
        ws['A3'].font = Font(bold=True)
        
        user_count = sum(1 for m in messages if m.get('role') == 'user')
        assistant_count = sum(1 for m in messages if m.get('role') == 'assistant')
        
        ws.append(['Gesamt Messages:', len(messages)])
        ws.append(['User Messages:', user_count])
        ws.append(['Assistant Messages:', assistant_count])
        ws.append([])
        
        # Performance Statistics
        ws.append(['Performance-Statistik'])
        ws['A8'].font = Font(bold=True)
        
        durations = [
            m.get('metadata', {}).get('duration', 0)
            for m in messages
            if m.get('role') == 'assistant' and m.get('metadata', {}).get('duration')
        ]
        
        if durations:
            ws.append(['Durchschnittliche Antwortzeit:', f"{sum(durations) / len(durations):.2f}s"])
            ws.append(['Schnellste Antwort:', f"{min(durations):.2f}s"])
            ws.append(['Langsamste Antwort:', f"{max(durations):.2f}s"])
        
        ws.append([])
        
        # Feedback Statistics
        if feedback_stats:
            ws.append(['Feedback-Statistik'])
            ws['A13'].font = Font(bold=True)
            
            ws.append(['Gesamt Feedback:', feedback_stats.get('total_feedback', 0)])
            ws.append(['Positiv:', feedback_stats.get('positive_count', 0)])
            ws.append(['Negativ:', feedback_stats.get('negative_count', 0)])
            ws.append(['Positive Ratio:', f"{feedback_stats.get('positive_ratio', 0):.1f}%"])
        
        # Column widths
        ws.column_dimensions['A'].width = 30
        ws.column_dimensions['B'].width = 20
    
    def _create_sources_sheet(self, ws, messages: List[Dict]):
        """Erstellt Quellen-Sheet"""
        # Header
        headers = ['Message Nr.', 'Source Title', 'Page', 'Relevance']
        ws.append(headers)
        
        # Style Header
        for cell in ws[1]:
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill(start_color="70AD47", end_color="70AD47", fill_type="solid")
        
        # Data
        for i, msg in enumerate(messages, 1):
            if msg.get('role') != 'assistant':
                continue
            
            sources = msg.get('sources', [])
            for source in sources:
                row = [
                    i,
                    source.get('title', 'Unbekannt'),
                    source.get('page', ''),
                    source.get('relevance', '')
                ]
                ws.append(row)
        
        # Column widths
        ws.column_dimensions['A'].width = 15
        ws.column_dimensions['B'].width = 50
        ws.column_dimensions['C'].width = 10
        ws.column_dimensions['D'].width = 15
    
    # ===== UTILITY =====
    
    def filter_messages_by_date(
        self,
        messages: List[Dict],
        days: int = 7
    ) -> List[Dict]:
        """
        Filtert Messages nach Zeitraum
        
        Args:
            messages: Alle Messages
            days: Anzahl Tage (default: 7)
            
        Returns:
            Gefilterte Messages
        """
        cutoff_date = datetime.now() - timedelta(days=days)
        
        filtered = []
        for msg in messages:
            timestamp_str = msg.get('timestamp', '')
            
            try:
                # Parse verschiedene Formate
                for fmt in ['%Y-%m-%d %H:%M:%S', '%d.%m.%Y %H:%M', '%Y-%m-%dT%H:%M:%S']:
                    try:
                        msg_date = datetime.strptime(timestamp_str, fmt)
                        if msg_date >= cutoff_date:
                            filtered.append(msg)
                        break
                    except ValueError:
                        continue
            except:
                # Bei Fehler: Message behalten
                filtered.append(msg)
        
        return filtered
    
    def get_supported_formats(self) -> List[str]:
        """Returns Liste unterstützter Formate"""
        formats = []
        if DOCX_AVAILABLE:
            formats.append('.docx')
        if OPENPYXL_AVAILABLE:
            formats.append('.xlsx')
        return formats
